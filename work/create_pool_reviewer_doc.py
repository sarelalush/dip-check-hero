from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "pool-chemistry-review-guide.docx"


COLORS = {
    "navy": "0B2545",
    "cyan": "11B5C8",
    "green": "20B86F",
    "green_light": "E8F8EF",
    "yellow": "F4C542",
    "yellow_light": "FFF7D6",
    "red": "E9576A",
    "red_light": "FFE6EA",
    "cyan_light": "E6F9FC",
    "gray": "EEF3F6",
    "gray_text": "516778",
    "border": "B7DCE4",
    "white": "FFFFFF",
}


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8E6EA", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def rtl_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    return paragraph


def format_run(run, size=11, bold=False, color="0B2545"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_rtl_paragraph(doc, text="", size=11, bold=False, color="0B2545", space_after=6):
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        format_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    size = {1: 17, 2: 14, 3: 12}.get(level, 12)
    color = COLORS["cyan"] if level == 1 else COLORS["navy"]
    p = add_rtl_paragraph(doc, text, size=size, bold=True, color=color, space_after=5)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    return p


def style_cell(cell, fill=None, text_color="0B2545", bold=False, size=9.5):
    if fill:
        set_cell_fill(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        rtl_paragraph(p)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            format_run(run, size=size, bold=bold, color=text_color)


def add_table(doc, headers, rows, widths, header_fill="E6F9FC"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    header_cells = table.rows[0].cells
    for i, value in enumerate(headers):
        if isinstance(value, tuple):
            text, fill, color, bold = value
        else:
            text, fill, color, bold = value, header_fill, COLORS["navy"], True
        header_cells[i].text = str(text)
        style_cell(header_cells[i], fill=fill, text_color=color, bold=bold, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, tuple):
                text, fill, color, bold = value
            else:
                text, fill, color, bold = value, None, COLORS["navy"], False
            cells[i].text = str(text)
            style_cell(cells[i], fill=fill, text_color=color, bold=bold, size=9.2)
    doc.add_paragraph()
    return table


def add_callout(doc, text, fill="FFF7D6", border="F4C542"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_fill(cell, fill)
    set_cell_border(cell, color=border, size="8")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        rtl_paragraph(p)
        for run in p.runs:
            format_run(run, size=10.5, bold=True, color=COLORS["navy"])
    doc.add_paragraph()


def text_color_for_fill(fill):
    red = int(fill[0:2], 16)
    green = int(fill[2:4], 16)
    blue = int(fill[4:6], 16)
    luminance = 0.299 * red + 0.587 * green + 0.114 * blue
    return COLORS["navy"] if luminance > 155 else COLORS["white"]


def add_color_scale(doc, title, rows, fills, note):
    add_heading(doc, title, 2)
    column_width = 5.55 / len(fills)
    table = doc.add_table(rows=0, cols=len(fills) + 1)
    table.autofit = False
    table.columns[0].width = Inches(1.05)
    for index in range(1, len(fills) + 1):
        table.columns[index].width = Inches(column_width)

    for label, values in rows:
        cells = table.add_row().cells
        cells[0].text = label
        style_cell(cells[0], fill=COLORS["cyan_light"], text_color=COLORS["navy"], bold=True, size=8.8)
        for index, (value, fill) in enumerate(zip(values, fills), start=1):
            cells[index].text = str(value)
            style_cell(
                cells[index],
                fill=fill,
                text_color=text_color_for_fill(fill),
                bold=True,
                size=9,
            )

    add_rtl_paragraph(doc, note, size=9.2, color=COLORS["gray_text"], space_after=7)


def add_parameter_behavior(doc, entries):
    rows = []
    for fill, color_name, value, status, action in entries:
        rows.append(
            [
                (color_name, fill, text_color_for_fill(fill), True),
                value,
                status,
                action,
            ]
        )
    add_table(
        doc,
        ["הצבע שרואים על הפד", "ערך", "משמעות", "מה מציעים למשתמש"],
        rows,
        [1.45, 0.8, 1.2, 2.8],
    )


def main():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    styles["Normal"].font.size = Pt(10.5)

    title = add_rtl_paragraph(
        doc,
        "AquaSense - מדריך לבדיקת החלטות צבע והמלצות טיפול",
        size=20,
        bold=True,
        color=COLORS["navy"],
        space_after=2,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = add_rtl_paragraph(
        doc,
        "מסמך לבודק בריכות: איך האפליקציה מחליטה מה תקין, מה חסר ומה צריך להוסיף",
        size=11,
        bold=True,
        color=COLORS["gray_text"],
        space_after=10,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_callout(
        doc,
        "חשוב: האפליקציה לא מטפלת לפי צבע הקובייה בלבד. קודם היא מזהה מהסטיק ערך מספרי, ואז משווה אותו לטווחים. הצבע בקובייה הוא תוצאה של ההשוואה הזו.",
        fill="E6F9FC",
        border=COLORS["cyan"],
    )

    add_heading(doc, "1. טבלת הצבעים הפיזית של AquaChek Pro", 1)
    add_callout(
        doc,
        "צבע הפד על הסטיק אינו צבע הסטטוס באפליקציה. לדוגמה: פד pH אדום עשוי לייצג 7.8 או 8.4, בעוד שפד כלור חופשי גבוה הוא סגול. רק לאחר המרת צבע הפד לערך מספרי האפליקציה קובעת נמוך, תקין או גבוה.",
        fill="F3F7F9",
        border=COLORS["border"],
    )
    add_table(
        doc,
        ["מיקום על הסטיק", "מה הפד מודד", "משפחת הצבעים"],
        [
            ["1 - קצה רטוב", "כלור כללי / ברום כללי באותו פד", "צהוב בהיר -> ירוק כהה"],
            ["2", "כלור חופשי", "צהוב-לבן -> סגול כהה"],
            ["3", "pH", "צהוב-כתום -> אדום כהה"],
            ["4 - קרוב לידית", "אלקליניות כללית", "צהוב -> ירוק -> כחול-ירקרק"],
        ],
        [1.35, 2.8, 2.1],
    )

    add_color_scale(
        doc,
        "פד משולב: כלור כללי וברום כללי (ppm)",
        [
            ("כלור כללי", ["0", "0.5", "1", "3", "5", "10"]),
            ("ברום כללי", ["0", "1", "2", "5", "10", "20"]),
        ],
        ["FFF9A8", "F2FEAA", "E7F5A0", "B8D88C", "64B469", "378C50"],
        "זהו פד פיזי אחד עם שתי סקאלות. בבריכת כלור קוראים את שורת הכלור; בבריכת ברום קוראים את שורת הברום.",
    )
    add_parameter_behavior(
        doc,
        [
            ("FFF9A8", "צהוב בהיר", "TC 0 / TB 0", "נמוך מאוד", "אין כמעט חומר חיטוי; לבדוק גם כלור חופשי ולהוסיף לפי סוג הבריכה."),
            ("F2FEAA", "צהוב-ירקרק", "TC 0.5 / TB 1", "נמוך", "רמת החיטוי נמוכה; להעלות בהדרגה."),
            ("E7F5A0", "ירוק בהיר", "TC 1 / TB 2", "קצה תחתון תקין", "לרוב אין תיקון מיידי; לוודא שכלור חופשי תקין."),
            ("B8D88C", "ירוק", "TC 3 / TB 5", "אידיאלי", "אין פעולה."),
            ("64B469", "ירוק בינוני", "TC 5 / TB 10", "קצה עליון", "לא להוסיף חומר כרגע; לבדוק שוב בהמשך."),
            ("378C50", "ירוק כהה", "TC 10 / TB 20", "גבוה", "להפסיק הוספה, להפעיל סירקולציה ולהמתין."),
        ],
    )
    add_color_scale(
        doc,
        "כלור חופשי (ppm)",
        [("ערך", ["0", "0.5", "1", "3", "5", "10"])],
        ["FEFECC", "F7EBE4", "EBE4E2", "A78BC7", "9C63B4", "812E9C"],
        "לבריכה: 0-0.5 נמוך; 1-3 הוא אזור העבודה; 5 ומעלה גבוה. שימו לב: כלור חופשי גבוה נראה סגול, לא אדום.",
    )
    add_parameter_behavior(
        doc,
        [
            ("FEFECC", "צהוב בהיר", "0", "נמוך מאוד", "להוסיף כלור לפי נפח הבריכה והוראות החומר, ואז לבדוק שוב."),
            ("F7EBE4", "לבן-קרם", "0.5", "נמוך", "להעלות כלור בהדרגה."),
            ("EBE4E2", "אפור-לילך בהיר", "1", "קצה תחתון תקין", "אין פעולה מיידית; לעקוב."),
            ("A78BC7", "סגול בהיר", "3", "קצה עליון תקין", "אין פעולה."),
            ("9C63B4", "סגול", "5", "גבוה לבריכה", "לא להוסיף כלור; להפעיל סירקולציה ולהמתין."),
            ("812E9C", "סגול כהה", "10", "גבוה מאוד", "להימנע מרחצה, לעצור הוספה ולפנות לאיש בריכות אם הערך אינו יורד."),
        ],
    )
    add_color_scale(
        doc,
        "pH",
        [("ערך", ["6.2", "6.8", "7.2", "7.8", "8.4"])],
        ["EFB52F", "E87522", "DC4A23", "D52F25", "BE292D"],
        "לפי מקרא היצרן: 6.2-6.8 נמוך, 7.2-7.8 מסומן OK, ו-8.4 גבוה. האפליקציה משתמשת כרגע ב-7.2-7.6 כטווח תקין - נדרש אישור מקצועי לפער.",
    )
    add_parameter_behavior(
        doc,
        [
            ("EFB52F", "צהוב-כתום", "6.2", "נמוך מאוד", "לבדוק קודם אלקליניות; לאחר מכן להעלות pH בהדרגה."),
            ("E87522", "כתום", "6.8", "נמוך", "להעלות pH בהדרגה ולבדוק שוב."),
            ("DC4A23", "כתום-אדום", "7.2", "תקין", "אין פעולה."),
            ("D52F25", "אדום", "7.8", "גבולי", "היצרן מסמן OK אך האפליקציה מחשיבה מעל 7.6 כגבוה; נדרש אישור הבודק."),
            ("BE292D", "אדום כהה", "8.4", "גבוה", "להוריד pH בהדרגה בהתאם לחומר ולנפח; מומלץ איש בריכות."),
        ],
    )
    add_color_scale(
        doc,
        "אלקליניות כללית (ppm)",
        [("ערך", ["0", "40", "80", "120", "180", "240"])],
        ["E3C040", "A4A933", "899F3A", "557F55", "376964", "285A78"],
        "0-40 נמוך, 80-120 תקין, 180-240 גבוה. בקצה הגבוה הפד כחול-ירקרק; כחול כאן אינו סימן לתוצאה טובה.",
    )
    add_parameter_behavior(
        doc,
        [
            ("E3C040", "צהוב", "0", "נמוך מאוד", "להוסיף מעלה אלקליניות לפי נפח הבריכה ובהדרגה."),
            ("A4A933", "צהוב-זית", "40", "נמוך", "להעלות אלקליניות ולבדוק שוב לפני תיקון pH."),
            ("899F3A", "ירוק-זית", "80", "קצה תחתון תקין", "אין פעולה."),
            ("557F55", "ירוק", "120", "קצה עליון תקין", "אין פעולה; לעקוב."),
            ("376964", "ירוק-כחול", "180", "גבוה", "להוריד בהדרגה לפי הנחיית איש בריכות; לא לבצע שינוי חד."),
            ("285A78", "כחול כהה", "240", "גבוה מאוד", "נדרש טיפול מקצועי והורדה מדורגת."),
        ],
    )

    add_callout(
        doc,
        "ל-AquaChek Pro שבתמונות אין פד נפרד של CYA / מייצב כלור. אין להסיק CYA מארבעת הפדים האלה. CYA יוצג רק אם מותג הסטיק שנבחר כולל מדידה כזו או אם המשתמש הזין בדיקה מתאימה.",
        fill=COLORS["red_light"],
        border=COLORS["red"],
    )

    add_heading(doc, "2. צבעי הממשק בלבד - לא צבעי הסטיק", 1)
    add_table(
        doc,
        ["צבע שרואים במסך", "משמעות", "מה האפליקציה חושבת", "מה קורה בפועל"],
        [
            [("ירוק", COLORS["green"], "FFFFFF", True), "הערך תקין", "המדידה בתוך הטווח הרצוי", "לא ממליצה להוסיף חומר"],
            [("צהוב", COLORS["yellow"], "0B2545", True), "הערך נמוך", "חסר חומר או הערך מתחת לטווח", "ממליצה להעלות את הערך"],
            [("אדום", COLORS["red"], "FFFFFF", True), "הערך גבוה", "יש עודף חומר או ערך מעל הטווח", "ממליצה להוריד או להמתין לפי סוג המדד"],
            [("טורקיז", COLORS["cyan"], "FFFFFF", True), "לטיפול עכשיו", "זו הפעולה הראשונה בסדר העדיפות", "מציגה מינון/פעולה לביצוע עכשיו"],
            [("אפור", COLORS["gray"], "0B2545", True), "אין נתון", "המדידה לא זוהתה או חסרה", "לא נותנת המלצה לפרמטר הזה"],
        ],
        [1.15, 1.1, 2.1, 2.15],
    )

    add_heading(doc, "3. הטווחים שעליהם מבוססת ההחלטה", 1)
    add_table(
        doc,
        ["קובייה", ("נמוך", COLORS["yellow_light"], COLORS["navy"], True), ("תקין", COLORS["green_light"], COLORS["navy"], True), ("גבוה", COLORS["red_light"], COLORS["navy"], True), "יעד"],
        [
            ["כלור חופשי", "פחות מ-1 ppm", "1-3 ppm", "מעל 3 ppm", "2 ppm"],
            ["pH", "פחות מ-7.2", "7.2-7.6", "מעל 7.6", "7.4"],
            ["אלקליניות", "פחות מ-80 ppm", "80-120 ppm", "מעל 120 ppm", "100 ppm"],
            ["מייצב כלור / CYA", "לא נמדד", "לא נמדד", "לא נמדד", "לא קיים בסטיק"],
        ],
        [1.45, 1.35, 1.25, 1.25, 0.9],
    )

    add_heading(doc, "4. סדר עדיפות שהאפליקציה מפעילה", 1)
    add_table(
        doc,
        ["עדיפות", "פרמטר", "למה קודם"],
        [
            ["1", "אלקליניות", "אם האלקליניות לא מאוזנת, קשה לייצב pH וכל טיפול אחר פחות מדויק."],
            ["2", "pH", "pH משפיע גם על נוחות הרחצה וגם על יעילות הכלור."],
            ["3", "כלור חופשי", "אחרי שהמים יציבים יותר, מתקנים את רמת החיטוי."],
            ["4", "CYA / מייצב כלור", "רק במותג סטיק שמודד CYA או לאחר הזנה ידנית; AquaChek Pro שבתמונה אינו מודד אותו."],
        ],
        [0.65, 1.5, 4.35],
    )

    add_heading(doc, "5. חריגים חשובים", 1)
    add_callout(
        doc,
        "אם pH נמוך מאוד מתחת ל-7.2, האפליקציה מתייחסת לזה כמצב בטיחותי. אם במקביל האלקליניות גבוהה, ההמלצה היא אוורור וסירקולציה ולא הוספת חומצה.",
        fill="FFF7D6",
        border=COLORS["yellow"],
    )
    add_callout(
        doc,
        "אם אין נפח בריכה בליטרים, האפליקציה יכולה להציג כיוון טיפול, אבל לא תחשב מינון מדויק.",
        fill="FFE6EA",
        border=COLORS["red"],
    )

    add_heading(doc, "6. מה לבקש מבודק הבריכות לאשר", 1)
    checklist = [
        "האם הטווחים 1-3 כלור חופשי, 7.2-7.6 pH ו-80-120 אלקליניות מתאימים לבריכות פרטיות בארץ?",
        "האם נכון לקרוא את הפד המשולב כ-TC 0/0.5/1/3/5/10 או TB 0/1/2/5/10/20 לפי סוג הבריכה?",
        "האם מוסכם שאין להציג CYA כתוצאה של AquaChek Pro, משום שאין בסטיק הזה פד CYA?",
        "האם סדר העדיפות אלקליניות -> pH -> כלור -> CYA נכון מקצועית?",
        "האם ההחלטה לטפל בפעולה אחת בכל פעם נכונה?",
        "האם ההמלצה במקרה pH נמוך + אלקליניות גבוהה לבצע אוורור במקום חומצה נכונה?",
        "האם ב-CYA גבוה נכון להמליץ על החלפת חלק מהמים?",
        "האם בכלור גבוה מעט נכון להמליץ על סירקולציה והמתנה לפני מנטרל כלור?",
    ]
    for item in checklist:
        p = add_rtl_paragraph(doc, f"☐ {item}", size=10.2, color=COLORS["navy"], space_after=4)
        p.paragraph_format.left_indent = Inches(0.1)

    add_heading(doc, "הערת אחריות", 1)
    add_rtl_paragraph(
        doc,
        "ההמלצות הן אינדיקציה בלבד לפי נתוני הסטיק ונפח הבריכה. יש לפעול לפי הוראות יצרני חומרי הבריכה, להוסיף חומרים בהדרגה, לא לערבב חומרים שונים יחד, ובמקרה של ספק להתייעץ עם איש בריכות מוסמך.",
        size=10.5,
        bold=True,
        color=COLORS["gray_text"],
        space_after=4,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
