from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "aquachek-pro-treatment-and-dosage-guide.docx"

NAVY = "0B2545"
CYAN = "11B5C8"
LIGHT_CYAN = "E6F9FC"
GRAY = "516778"
WHITE = "FFFFFF"
BORDER = "B7DCE4"


def rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def format_run(run, size=10, bold=False, color=NAVY):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", size=10, bold=False, color=NAVY, center=False, after=5):
    p = doc.add_paragraph()
    rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        format_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, size=17):
    p = paragraph(doc, text, size=size, bold=True, color=CYAN, after=6)
    p.paragraph_format.space_before = Pt(8)
    return p


def set_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_border(cell, color="D8E6EA", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def style_cell(cell, fill=None, text_color=NAVY, bold=False, size=9):
    if fill:
        set_fill(cell, fill)
    set_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        rtl(p)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            format_run(run, size=size, bold=bold, color=text_color)


def text_color(fill):
    red, green, blue = (int(fill[i : i + 2], 16) for i in (0, 2, 4))
    luminance = 0.299 * red + 0.587 * green + 0.114 * blue
    return NAVY if luminance > 155 else WHITE


def callout(doc, text, fill="F3F7F9", border=BORDER):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    set_fill(cell, fill)
    set_border(cell, border, "8")
    style_cell(cell, fill=fill, text_color=NAVY, bold=True, size=10)
    paragraph(doc, after=3)


def indicator_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [1.25, 0.8, 1.3, 3.0]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)

    headers = ["צבע הפד בפועל", "ערך", "פירוש", "המלצה לבדיקה מקצועית"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
        style_cell(cell, fill=LIGHT_CYAN, text_color=NAVY, bold=True, size=9)

    for fill, color_name, value, meaning, action in rows:
        cells = table.add_row().cells
        cells[0].text = color_name
        style_cell(cells[0], fill=fill, text_color=text_color(fill), bold=True, size=9)
        for cell, value_text in zip(cells[1:], (value, meaning, action)):
            cell.text = value_text
            style_cell(cell, size=9)
    paragraph(doc, after=2)


def dosage_table(doc, rows):
    paragraph(doc, "מינון שהאפליקציה מחשבת כיום - בסיס של 10,000 ליטר", size=11, bold=True, color=CYAN, after=4)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [0.9, 2.0, 2.25, 1.75]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)

    headers = ["ערך", "מה מוסיפים", "כמות ל-10,000 ליטר", "הערה מקצועית"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
        style_cell(cell, fill=LIGHT_CYAN, text_color=NAVY, bold=True, size=8.5)

    for value, product, amount, note in rows:
        cells = table.add_row().cells
        for cell, value_text in zip(cells, (value, product, amount, note)):
            cell.text = value_text
            style_cell(cell, size=8.5)
    paragraph(doc, after=2)


def start_indicator(doc, number, title, subtitle):
    if number > 1:
        doc.add_page_break()
    paragraph(doc, f"בדיקה {number}", size=10, bold=True, color=GRAY, after=1)
    heading(doc, title, 19)
    paragraph(doc, subtitle, size=10.5, bold=True, color=GRAY, after=8)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    doc.styles["Normal"].font.size = Pt(10)

    paragraph(doc, "AquaSense", size=14, bold=True, color=CYAN, center=True, after=2)
    paragraph(doc, "מדריך מקצועי לפענוח AquaChek Pro", size=21, bold=True, center=True, after=3)
    paragraph(
        doc,
        "כל אינדיקציה נבדקת בנפרד לפי משפחת הצבעים שלה. אין צבע אוניברסלי שמשמעותו תקין או לא תקין.",
        size=11,
        bold=True,
        color=GRAY,
        center=True,
        after=10,
    )
    callout(
        doc,
        "דוגמה: סגול כהה בכלור חופשי מצביע על ערך גבוה; אדום כהה ב-pH מצביע על ערך גבוה; ובאלקליניות אין אדום כלל - הערך הגבוה הוא ירוק-כחול או כחול כהה.",
        fill=LIGHT_CYAN,
        border=CYAN,
    )
    paragraph(doc, "סדר ארבעת הפדים מהקצה הרטוב לכיוון הידית:", size=11, bold=True)
    paragraph(doc, "1. כלור כללי / ברום כללי  |  2. כלור חופשי  |  3. pH  |  4. אלקליניות כללית", size=10.5)
    callout(
        doc,
        "בסטיק הזה אין פד CYA / מייצב כלור. אין להפיק ממנו תוצאת CYA.",
        fill="FFF1F3",
        border="E9576A",
    )
    callout(
        doc,
        "כל כמויות הטיפול במסמך מחושבות לבריכה של 10,000 ליטר. להתאמה לנפח אחר: הכמות בפועל = הכמות בטבלה × נפח הבריכה בליטרים ÷ 10,000. החישובים משקפים את קוד האפליקציה הנוכחי ומיועדים לאישור איש הבריכות לפני פרסום כהנחיה סופית.",
        fill="FFF8E8",
        border="E7B84B",
    )
    paragraph(doc, "הנחות החישוב הנוכחיות: כלור נוזלי 12%, חומצת מלח 33%, יעד כלור חופשי 2 ppm, יעד pH ‏7.4 ויעד אלקליניות 100 ppm.", size=9.5, bold=True, color=GRAY, after=7)

    start_indicator(
        doc,
        1,
        "כלור כללי / ברום כללי (ppm)",
        "זהו פד משולב אחד. משתמשים בסקאלת הכלור בבריכת כלור ובסקאלת הברום בבריכת ברום.",
    )
    indicator_table(
        doc,
        [
            ("FFF9A8", "צהוב בהיר", "TC 0 / TB 0", "נמוך מאוד", "אין כמעט חומר חיטוי. לאשר מול כלור חופשי ולהעלות חומר בהתאם לסוג הבריכה."),
            ("F2FEAA", "צהוב-ירקרק", "TC 0.5 / TB 1", "נמוך", "רמת החיטוי נמוכה; להעלות בהדרגה ולבדוק שוב."),
            ("E7F5A0", "ירוק בהיר", "TC 1 / TB 2", "קצה תחתון", "להצליב עם כלור חופשי לפני המלצה."),
            ("B8D88C", "ירוק", "TC 3 / TB 5", "אידיאלי", "אין פעולה לפי הפד הזה בלבד."),
            ("64B469", "ירוק בינוני", "TC 5 / TB 10", "קצה עליון", "לא להוסיף כרגע; להצליב עם כלור חופשי."),
            ("378C50", "ירוק כהה", "TC 10 / TB 20", "גבוה", "לעצור הוספה, להפעיל סירקולציה ולבדוק שוב."),
            ("28663E", "ירוק כהה מאוד", "TC 20 / TB 40", "גבוה מאוד", "לעצור הוספה; לאפשר רחצה רק לאחר בדיקה מקצועית וירידת הערך."),
        ],
    )
    dosage_table(
        doc,
        [
            ("TC 0-20", "אין חומר לפי פד זה בלבד", "אין מינון ישיר", "בבריכת כלור המינון נקבע לפי כלור חופשי והפער TC-FC."),
            ("TB 0-40", "טרם הוגדר חומר ברום", "אין מינון באפליקציה", "נדרש להגדיר מוצר ברום, ריכוז ויעד לפני חישוב."),
        ],
    )
    callout(
        doc,
        "הפד המשולב אינו זהה לכלור חופשי. בבריכת כלור אפשר לחשב כלור משולב בקירוב: כלור כללי פחות כלור חופשי. פער גדול דורש בדיקה מקצועית.",
    )

    start_indicator(
        doc,
        2,
        "כלור חופשי (ppm)",
        "משפחת הצבעים: צהוב-לבן בערכים נמוכים, וסגול בערכים הגבוהים. כאן סגול כהה הוא לא תקין.",
    )
    indicator_table(
        doc,
        [
            ("FEFECC", "צהוב בהיר", "0", "נמוך מאוד", "להוסיף כלור לפי נפח הבריכה והוראות היצרן, ואז לבדוק שוב."),
            ("F7EBE4", "לבן-קרם", "0.5", "נמוך", "להעלות כלור בהדרגה."),
            ("EBE4E2", "אפור-לילך בהיר", "1", "קצה תחתון תקין", "לרוב אין פעולה מיידית; לעקוב."),
            ("A78BC7", "סגול בהיר", "3", "קצה עליון תקין", "אין פעולה."),
            ("9C63B4", "סגול", "5", "גבוה לבריכה", "להפסיק הוספת כלור, להפעיל סירקולציה ולבדוק שוב."),
            ("812E9C", "סגול כהה", "10", "גבוה מאוד", "לא לאפשר רחצה; לעצור הוספה ולפנות לאיש בריכות אם הערך אינו יורד."),
            ("4E155F", "סגול כהה מאוד", "20", "קיצוני", "לא לאפשר רחצה; לעצור הוספה ולקבל הנחיה מקצועית."),
        ],
    )
    dosage_table(
        doc,
        [
            ("0 ppm", "כלור נוזלי 12%", "200 מ״ל", "מעלה בקירוב ליעד 2 ppm; בהנחה שאין טבליה פעילה."),
            ("0.5 ppm", "כלור נוזלי 12%", "150 מ״ל", "בהנחה שאין טבליה פעילה."),
            ("1 ppm", "לא מוסיפים", "0", "נמצא בגבול התחתון של הטווח 1-3 ppm."),
            ("3 ppm", "לא מוסיפים", "0", "נמצא בגבול העליון של הטווח."),
            ("5 ppm", "אנטי-כלור", "30 גרם", "המוצר והריכוז טרם הוגדרו; חובה לאשר לפני שימוש."),
            ("10 ppm", "אנטי-כלור", "80 גרם", "ערך גבוה מאוד; חובה לאשר מוצר ומינון."),
            ("20 ppm", "אנטי-כלור", "180 גרם", "ערך קיצוני; טיפול מקצועי בלבד."),
        ],
    )
    callout(
        doc,
        "כאשר מוגדרת טבליית כלור פעילה, האפליקציה מפחיתה מהכלור הנוזלי את תרומת הטבליה לפי משקל, כמות, שעות משאבה וזמן ההמתנה. לכן 200/150 מ״ל הם ערכי בסיס ללא טבליה פעילה.",
    )

    start_indicator(
        doc,
        3,
        "pH",
        "משפחת הצבעים: צהוב-כתום בערכים נמוכים ואדום בערכים הגבוהים. כאן אדום כהה מצביע על pH גבוה.",
    )
    indicator_table(
        doc,
        [
            ("EFB52F", "צהוב-כתום", "6.2", "נמוך מאוד", "לבדוק אלקליניות תחילה; לאחר מכן להעלות pH בהדרגה."),
            ("E87522", "כתום", "6.8", "נמוך", "להעלות pH בהדרגה ולבדוק שוב."),
            ("DC4A23", "כתום-אדום", "7.2", "תקין", "אין פעולה."),
            ("D52F25", "אדום", "7.8", "גבולי", "היצרן מסמן OK, אך האפליקציה משתמשת כרגע בגבול 7.6. נדרש אישור הבודק."),
            ("BE292D", "אדום כהה", "8.4", "גבוה", "להוריד pH בהדרגה ובהתאם לנפח; מומלץ איש בריכות."),
        ],
    )
    dosage_table(
        doc,
        [
            ("6.2", "pH Plus", "1,200 גרם", "החישוב הנוכחי גדול ודורש אימות מול מוצר ואלקליניות."),
            ("6.8", "pH Plus", "600 גרם", "להוסיף בהדרגה, לסחרר ולבדוק שוב."),
            ("7.2", "לא מוסיפים", "0", "בתוך הטווח הנוכחי 7.2-7.6."),
            ("7.8", "חומצת מלח 33%", "200 מ״ל", "האפליקציה מסווגת מעל 7.6 כגבוה; נדרש אישור יעד."),
            ("8.4", "חומצת מלח 33%", "500 מ״ל", "ערך קיצוני; להוסיף בהדרגה ורק לפי הנחיה מקצועית."),
        ],
    )
    callout(
        doc,
        "חריג בטיחות בקוד: אם pH נמוך מ-7.2 והאלקליניות גבוהה, האפליקציה אינה מוסיפה חומצה או pH Plus אלא ממליצה על אוורור, סחרור ובדיקה חוזרת.",
    )

    start_indicator(
        doc,
        4,
        "אלקליניות כללית (ppm)",
        "משפחת הצבעים: צהוב בערכים נמוכים, ירוק בטווח העבודה וכחול-ירקרק בערכים הגבוהים. אין אדום בסקאלה הזו.",
    )
    indicator_table(
        doc,
        [
            ("E3C040", "צהוב", "0", "נמוך מאוד", "להוסיף מעלה אלקליניות לפי נפח הבריכה ובהדרגה."),
            ("A4A933", "צהוב-זית", "40", "נמוך", "להעלות אלקליניות ולבדוק שוב לפני תיקון pH."),
            ("899F3A", "ירוק-זית", "80", "קצה תחתון תקין", "אין פעולה."),
            ("557F55", "ירוק", "120", "קצה עליון תקין", "אין פעולה; לעקוב."),
            ("376964", "ירוק-כחול", "180", "גבוה", "להוריד בהדרגה לפי הנחיית איש בריכות; לא לבצע שינוי חד."),
            ("285A78", "כחול כהה", "240", "גבוה מאוד", "נדרש טיפול מקצועי והורדה מדורגת."),
        ],
    )
    dosage_table(
        doc,
        [
            ("0 ppm", "סודה לשתייה / Alkalinity Increaser", "1.5 ק״ג", "מכוון ל-100 ppm לפי הנוסחה הנוכחית."),
            ("40 ppm", "סודה לשתייה / Alkalinity Increaser", "0.9 ק״ג", "מכוון ל-100 ppm לפי הנוסחה הנוכחית."),
            ("80 ppm", "לא מוסיפים", "0", "בתוך הטווח 80-120 ppm."),
            ("120 ppm", "לא מוסיפים", "0", "בתוך הטווח; מעקב בלבד."),
            ("180 ppm", "חומצת מלח 33%", "0.5 ליטר בשלב ראשון", "האומדן הכולל 1.6 ל׳; הקוד נותן שליש בכל טיפול."),
            ("240 ppm", "חומצת מלח 33%", "0.9 ליטר בשלב ראשון", "האומדן הכולל 2.8 ל׳; הקוד נותן שליש בכל טיפול."),
        ],
    )
    callout(
        doc,
        "בהורדת אלקליניות האפליקציה מציגה רק שליש מהאומדן בכל סבב. לאחר ההוספה: סחרור, המתנה של 4-6 שעות ובדיקה חוזרת לפני מנה נוספת.",
    )

    doc.add_page_break()
    heading(doc, "סדר עדיפויות מוצע", 19)
    paragraph(
        doc,
        "סדר העדיפויות נקבע לפי סוג המדד והערך המספרי, ולא לפי צבע הפד. זהו הסדר המוצע לבדיקת איש הבריכות:",
        size=10.5,
        bold=True,
        color=GRAY,
        after=8,
    )
    priorities = [
        ("חריג בטיחות", "כלור חופשי 10 ppm או pH קיצוני 6.2/8.4", "עוצר המלצות רגילות ומציג אזהרת רחצה וטיפול מקצועי."),
        ("1", "אלקליניות", "מאזנים קודם משום שהיא משפיעה על יציבות ה-pH."),
        ("2", "pH", "מתקנים לאחר האלקליניות משום שהוא משפיע על יעילות החיטוי, קורוזיה ואבנית."),
        ("3", "כלור חופשי", "לאחר שהמים יציבים יותר, מתקנים את רמת חומר החיטוי."),
        ("4", "כלור כללי / ברום כללי", "משמש להצלבה ולאבחון; בבריכת כלור בודקים גם את הפער מול כלור חופשי."),
    ]
    table = doc.add_table(rows=1, cols=3)
    for cell, label in zip(table.rows[0].cells, ("עדיפות", "בדיקה", "למה")):
        cell.text = label
        style_cell(cell, fill=LIGHT_CYAN, bold=True)
    for priority, test, reason in priorities:
        cells = table.add_row().cells
        for cell, value in zip(cells, (priority, test, reason)):
            cell.text = value
            style_cell(cell)
    paragraph(doc, after=3)

    callout(
        doc,
        "האפליקציה מציגה טיפול פעיל אחד בלבד בכל סריקה. היא אינה מורה להוסיף יחד סודה, חומצה וכלור. לאחר כל טיפול יש להפעיל סחרור, להמתין 4-6 שעות ולסרוק מחדש; רק אז נבדקת העדיפות הבאה.",
        fill=LIGHT_CYAN,
        border=CYAN,
    )

    heading(doc, "נוסחאות המינון הנוכחיות", 17)
    formulas = [
        "כלור חופשי נמוך: כל 1 ppm שחסר עד יעד 2 ppm = כ-100 מ״ל כלור נוזלי 12% לכל 10,000 ליטר, לפני זיכוי טבליה פעילה.",
        "pH נמוך: גרם pH Plus = (7.4 פחות pH שנמדד) × נפח הבריכה במ״ק × 100.",
        "pH גבוה: מ״ל חומצת מלח 33% = (pH שנמדד פחות 7.4) × נפח הבריכה במ״ק ÷ 20 × 1,000.",
        "אלקליניות נמוכה: ק״ג סודה לשתייה = (100 פחות ppm שנמדד) × נפח הבריכה במ״ק ÷ 670.",
        "אלקליניות גבוהה: אומדן ליטר חומצת מלח 33% = (ppm שנמדד פחות 100) × נפח הבריכה במ״ק ÷ 500; בכל סבב מוצג שליש בלבד.",
        "כלור גבוה מעל 4 ppm: גרם אנטי-כלור = (ppm שנמדד פחות 2) × נפח הבריכה במ״ק. הנוסחה אינה מאושרת עד שיוגדר מוצר וריכוז.",
    ]
    for item in formulas:
        paragraph(doc, f"• {item}", size=9.5, after=3)

    heading(doc, "פערים שחייבים אישור לפני פרסום", 17)
    gaps = [
        "לאנטי-כלור אין כרגע מוצר, חומר פעיל או ריכוז מוגדר. אין להציג את הכמות כהוראה סופית לפני אישור.",
        "מינון pH תלוי גם באלקליניות, בדרישת החומצה ובמוצר המסחרי; יש לאשר את נוסחאות 6.2/6.8/7.8/8.4.",
        "במובייל החישוב מציג כלור 12% וחומצת מלח 33%, בעוד שבקונפיגורציית הווב הישנה מופיעים 10% ו-32%. צריך לאחד את הריכוזים.",
        "אין כיום חישוב מינון לברום ואין מינון ישיר מכלור כללי; יש להחליט אם להוסיף תמיכה או להציג אבחון בלבד.",
        "יש לאשר שמנה ראשונה של שליש מאומדן החומצה היא שיטת העבודה הנכונה להורדת אלקליניות.",
    ]
    for item in gaps:
        paragraph(doc, f"☐ {item}", size=9.8, after=4)

    heading(doc, "נקודות לאישור בודק הבריכות", 17)
    checklist = [
        "האם כל ערכי הצבע והמעברים תואמים לבקבוק AquaChek Pro שבשימוש?",
        "האם 1-3 ppm הוא הטווח הנכון לכלור חופשי בבריכה פרטית?",
        "האם 7.8 צריך להיות תקין לפי היצרן או גבוה לפי יעד האפליקציה 7.2-7.6?",
        "האם 80-120 ppm הוא הטווח הנכון לאלקליניות כללית?",
        "האם סדר העדיפויות אלקליניות -> pH -> כלור חופשי נכון, ומהם החריגים?",
        "האם נכון להשתמש בכלור כללי פחות כלור חופשי כאינדיקציה לכלור משולב?",
        "האם מוסכם שאין להציג CYA כתוצאה של הסטיק הזה?",
        "האם כל חומר, ריכוז ומינון בטבלאות מתאים למוצרים שבפועל יומלצו למשתמשים?",
        "האם יש לאשר, לשנות או להסיר את מינוני האנטי-כלור עד שיוגדר מוצר מסוים?",
    ]
    for item in checklist:
        paragraph(doc, f"☐ {item}", size=10.2, after=4)

    heading(doc, "הערת אחריות", 17)
    paragraph(
        doc,
        "המסמך מיועד לאימות מקצועי של הלוגיקה ואינו הוראת טיפול סופית. מינון תלוי בנפח הבריכה, בסוג החומר ובריכוזו. יש לפעול לפי הוראות היצרן, לא לערבב חומרים, ולהתייעץ עם איש בריכות מוסמך במקרה של ספק או ערך קיצוני.",
        size=10.5,
        bold=True,
        color=GRAY,
    )

    heading(doc, "מקורות ייחוס", 15)
    sources = [
        "AquaChek Pro: https://www.aquachek.com/product/aquachek-pro/",
        "AquaChek FAQ - ההבדל בין כלור חופשי, כללי ומשולב: https://www.aquachek.com/faq/",
        "CDC - בדיקות וטיפול במי בריכות ביתיות: https://www.cdc.gov/healthy-swimming/about/home-pool-and-hot-tub-water-treatment-and-testing.html",
        "מקור נוסחאות האפליקציה: mobile/src/domain/dosage.ts והיעדים ב-mobile/src/config/targetRanges.ts.",
    ]
    for item in sources:
        paragraph(doc, item, size=8.5, color=GRAY, after=2)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
